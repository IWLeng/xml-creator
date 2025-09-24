import os
import re
import pandas as pd
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Alignment
from docx.enum.text import WD_COLOR_INDEX

def clean_text(text):
    """Apply regex replacements to clean the extracted text."""
    text = re.sub(r"\.wav", "", text)  # Remove ".wav"
    text = re.sub(r"content key.*", "", text, flags=re.IGNORECASE)  # Remove "content key:" and everything after
    text = re.sub(r"\<.*?\>", "", text)
    
    # Replace carriage returns with line feeds (Excel's Alt+Enter)
    text = text.replace('\r', '\n')
    
    return text.strip()

def extract_highlighted_text(cell):
    """Extract only YELLOW highlighted text from a DOCX table cell, preserving paragraph structure."""
    highlighted_text = []
    non_highlighted_text = []
    
    for para_idx, para in enumerate(cell.paragraphs):
        highlighted_runs = []
        non_highlighted_runs = []
        
        for run in para.runs:
            # Check if the text is highlighted with YELLOW specifically
            if run.font.highlight_color and run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                highlighted_runs.append(run.text)
            else:
                non_highlighted_runs.append(run.text)
        
        # Join runs within the same paragraph
        if highlighted_runs:
            highlighted_text.append(''.join(highlighted_runs))
        if non_highlighted_runs:
            non_highlighted_text.append(''.join(non_highlighted_runs))
        
        # Add line break between paragraphs (except after the last one)
        if para_idx < len(cell.paragraphs) - 1:
            if highlighted_text:  # Only add line break if we have highlighted content
                highlighted_text.append('\n')
            if non_highlighted_text:  # Only add line break if we have non-highlighted content
                non_highlighted_text.append('\n')
    
    highlighted = clean_text(''.join(highlighted_text))
    non_highlighted = clean_text(''.join(non_highlighted_text))
    
    return highlighted, non_highlighted

def autofit_columns(xlsx_path):
    """Auto-fit the columns in the Excel file with custom settings."""
    wb = load_workbook(xlsx_path)
    ws = wb.active

    # Set column B width to 130 and enable wrap text
    ws.column_dimensions['B'].width = 130
    
    # Apply wrap text ONLY to column B
    for row in ws.iter_rows(min_col=2, max_col=2):  # Column B only
        for cell in row:
            cell.alignment = Alignment(wrap_text=True)
    
    # Auto-fit column A (no wrap text)
    max_length = 0
    for cell in ws['A']:  # All cells in column A
        if cell.value:
            max_length = max(max_length, len(str(cell.value)))
    ws.column_dimensions['A'].width = max_length + 2  # Add some padding

    # Auto-fit other columns (C, D, etc. if they exist)
    for col in ws.columns:
        col_letter = col[0].column_letter
        if col_letter not in ['A', 'B']:  # Skip columns A and B since we already set them
            max_length = 0
            for cell in col:
                if cell.value:
                    # For cells with newlines, find the longest line
                    cell_lines = str(cell.value).split('\n')
                    for line in cell_lines:
                        max_length = max(max_length, len(line))
            adjusted_width = min(max_length + 2, 50)  # Cap at 50 to avoid extremely wide columns
            ws.column_dimensions[col_letter].width = adjusted_width

    wb.save(xlsx_path)

def process_docx_to_xlsx(docx_path):
    """Extract highlighted and non-highlighted text from DOCX and save to XLSX."""
    doc = Document(docx_path)
    extracted_data = []

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) > 0:  # Ensure at least one column exists
                # Try column 2 first (index 1)
                if len(row.cells) > 1:
                    cell = row.cells[1]  # Column 2
                    highlighted, non_highlighted = extract_highlighted_text(cell)
                    
                    # If no highlighted text in column 2, try column 3 (index 2)
                    if not highlighted and len(row.cells) > 2:
                        cell = row.cells[2]  # Column 3
                        highlighted, non_highlighted = extract_highlighted_text(cell)
                else:
                    # Only 1 column available, use column 1
                    cell = row.cells[0]  # Column 1
                    highlighted, non_highlighted = extract_highlighted_text(cell)
                
                if highlighted:  # Process only if there is highlighted text
                    extracted_data.append([highlighted, non_highlighted])

    if extracted_data:
        df = pd.DataFrame(extracted_data)
        xlsx_path = os.path.splitext(docx_path)[0] + ".xlsx"
        df.to_excel(xlsx_path, index=False, header=False)

        autofit_columns(xlsx_path)  # Auto-fit columns with custom settings
        print(f"Saved: {xlsx_path}")
    else:
        print(f"No YELLOW highlighted text found in: {docx_path}")

# Prompt user for folder path input
print("DOCX Preprocess for Icario")

folder_path = input("Enter the folder path containing DOCX files: ").strip()

if not os.path.isdir(folder_path):
    print("Invalid folder path. Exiting.")
else:
    docx_files = [f for f in os.listdir(folder_path) if f.endswith(".docx")]

    if not docx_files:
        print("No DOCX files found in the specified folder.")
    else:
        for docx_file in docx_files:
            docx_path = os.path.join(folder_path, docx_file)
            process_docx_to_xlsx(docx_path)